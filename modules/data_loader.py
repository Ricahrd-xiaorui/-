import streamlit as st
import os
import zipfile
import tempfile
import random
import string
import time
import re
import pandas as pd
import gc
from pathlib import Path
from utils.session_state import get_session_state, log_message, update_progress

# 尝试导入python-docx，如果失败则设置标志
try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✓ python-docx 导入成功")
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"✗ python-docx 导入失败: {e}")

# 尝试导入PyPDF2或pdfplumber，如果失败则设置标志
PDF_LIBRARY = None
try:
    import pdfplumber
    PDF_LIBRARY = "pdfplumber"
    print("✓ pdfplumber 导入成功")
except ImportError:
    try:
        import PyPDF2
        PDF_LIBRARY = "PyPDF2"
        print("✓ PyPDF2 导入成功")
    except ImportError as e:
        PDF_LIBRARY = None
        print(f"✗ PDF库导入失败: {e}")

# 大规模文件处理配置
MAX_FILES = 10000  # 最大支持文件数
BATCH_SIZE = 100   # 批量处理大小
MEMORY_WARNING_THRESHOLD = 1000  # 超过此数量显示内存警告

def load_example_data():
    """加载示例政策数据"""
    example_files = {
        "政策1_科技创新政策.txt": """
为全面贯彻党的十九大和十九届二中、三中、四中、五中全会精神，落实《国家创新驱动发展战略纲要》，推动科技创新与产业发展深度融合，加快建设科技强国，现提出如下意见。
一、总体要求
（一）指导思想。以习近平新时代中国特色社会主义思想为指导，深入贯彻党的十九大和十九届二中、三中、四中、五中全会精神，坚持创新在我国现代化建设全局中的核心地位，把科技自立自强作为国家发展的战略支撑，深入实施创新驱动发展战略，完善国家创新体系，加快建设科技强国，实现高水平科技自立自强。
（二）主要目标。到2025年，建成一批各具特色、优势互补、结构合理的科技创新基地，形成区域创新高地和特色创新集群，提升区域科技创新能力。加快推进科技成果转化应用，促进产业链创新链深度融合，构建现代产业技术体系，大幅提升企业创新能力，支撑引领高质量发展。
二、重点任务
（一）加强原创性引领性科技攻关。瞄准世界科技前沿、经济主战场、国家重大需求、人民生命健康，加强基础研究、注重原始创新，优化学科布局和研发布局，加强基础前沿交叉研究，促进学科交叉融合。
（二）促进科技创新与产业发展深度融合。完善产学研协同创新体系，大力推进科技成果转化应用，提升产业链现代化水平。加快数字化智能化绿色化转型，加强共性技术平台建设，促进中小企业创新发展，推动重点产业集群发展。
（三）培养造就高水平科技人才队伍。强化战略科技人才力量，培养具有国际水平的战略科技人才、科技领军人才、青年科技人才和高水平创新团队。完善人才发现、培养、激励机制，创新人才评价机制，构建具有国际竞争力的引才用才制度体系。
三、保障措施
（一）加大科技创新投入。完善政府投入为主、社会多渠道投入机制，促进科技投入持续稳定增长。引导企业加大研发投入，鼓励社会力量投入基础研究，建立健全基础研究多元投入机制。
（二）完善科技创新体制机制。深化科技体制改革，健全社会主义市场经济条件下新型举国体制，充分发挥政府作用和市场机制作用，完善科技治理体系，优化科技资源配置，促进科技事业全面发展。
（三）强化知识产权保护和运用。健全知识产权保护法律法规，完善知识产权保护体系，加强知识产权保护执法。健全知识产权运营服务体系，培育知识产权密集型产业，提升知识产权公共服务水平。
（四）深化科技开放合作。主动融入全球创新网络，坚持"引进来"和"走出去"相结合，促进国际科技交流合作，提高科技开放合作水平。优化创新创业环境，推动形成开放、融合、共享的创新生态。
本意见自发布之日起施行。
        """,
        "政策2_乡村振兴战略.txt": """
为深入贯彻党的十九大和十九届二中、三中、四中、五中全会精神，落实《中共中央国务院关于实施乡村振兴战略的意见》，全面推进乡村振兴战略实施，加快农业农村现代化，现提出如下意见。
一、总体要求
（一）指导思想。以习近平新时代中国特色社会主义思想为指导，深入贯彻党的十九大和十九届二中、三中、四中、五中全会精神，坚持农业农村优先发展，按照产业兴旺、生态宜居、乡风文明、治理有效、生活富裕的总要求，建立健全城乡融合发展体制机制和政策体系，加快推进农业农村现代化。
（二）主要目标。到2025年，乡村振兴取得重要进展，农业基础更加稳固，农村基础设施和公共服务体系基本健全，城乡融合发展体制机制基本建立。到2035年，乡村振兴取得决定性进展，农业农村现代化基本实现。到2050年，乡村全面振兴，农业强、农村美、农民富全面实现。
二、重点任务
（一）实施乡村产业振兴行动。大力发展乡村产业，做强做优农业，培育新产业新业态，拓展农业多种功能，推动产业融合发展，构建现代乡村产业体系，推动农村一二三产业融合发展。
（二）实施乡村人才振兴行动。加强农村人才队伍建设，培育新型职业农民，加强农村专业人才队伍建设，发挥科技人才支撑作用，鼓励社会各界人才投身乡村建设，创新乡村人才培育引进使用机制。
（三）实施乡村文化振兴行动。加强农村思想道德建设，弘扬和践行社会主义核心价值观，传承发展优秀传统文化，加强农村公共文化建设，开展移风易俗行动，推动形成文明乡风、良好家风、淳朴民风。
（四）实施乡村生态振兴行动。加强农村生态环境保护和治理，推动农业绿色发展，开展农村人居环境整治行动，加强农村生态保护与修复，健全农村生态文明制度体系，建设生态宜居美丽乡村。
（五）实施乡村组织振兴行动。加强农村基层党组织建设，健全乡村组织体系，完善村民自治机制，健全村级议事协商制度，发挥农村基层党组织战斗堡垒作用，增强乡村治理能力。
（六）实施乡村生活富裕行动。持续增加农民收入，优先发展农村教育事业，推进健康乡村建设，健全农村社会保障体系，改善农村人居环境，提高农村生活质量，促进农民生活富裕。
三、保障措施
（一）加强党对乡村振兴的领导。坚持党管农村工作，健全党委统一领导、政府负责、党委农村工作部门统筹协调的农村工作领导体制，建立健全实施乡村振兴战略领导责任制。
（二）完善乡村振兴政策体系。健全城乡融合发展体制机制，完善农业农村支持保护制度，加大乡村振兴投入力度，强化乡村振兴人才支撑。
（三）强化乡村振兴法治保障。完善乡村振兴法律法规体系，增强乡村法治观念，加强农村法治宣传教育，保障农民合法权益，推进法治乡村建设。
（四）加强乡村振兴考核评价。建立健全乡村振兴监测评价机制，开展乡村振兴战略实施监测评价，将乡村振兴战略实施成效纳入党政领导班子和领导干部实绩考核内容。
本意见自发布之日起施行。
        """,
        "政策3_数字经济发展政策.txt": """
为深入贯彻党的十九大和十九届二中、三中、四中、五中全会精神，落实《国家数字经济发展战略》，加快数字经济发展，推动数字技术与实体经济深度融合，现提出如下意见。
一、总体要求
（一）指导思想。以习近平新时代中国特色社会主义思想为指导，深入贯彻党的十九大和十九届二中、三中、四中、五中全会精神，立足新发展阶段，贯彻新发展理念，构建新发展格局，推动数字经济和实体经济深度融合，赋能传统产业转型升级，培育新产业新业态新模式，构建数字经济新优势。
（二）主要目标。到2025年，数字经济迈向全面扩展期，数字经济核心产业增加值占GDP比重达到10%，数字化创新引领发展能力大幅提升，智能化水平明显增强，数字技术与实体经济深度融合取得显著成效。到2035年，数字经济整体竞争力位居全球前列，数字化、网络化、智能化全面发展。
二、重点任务
（一）加快数字技术创新突破。强化基础研究和原始创新，加快人工智能、量子信息、集成电路、区块链、6G、物联网等关键数字技术攻关，推动数字技术创新突破。
（二）加速数字产业化发展。做大做强数字产业，培育壮大人工智能、大数据、云计算、区块链等新兴数字产业，加快推动数字产业集群发展，培育具有国际竞争力的数字企业和产业生态。
（三）推动产业数字化转型。加快传统产业数字化、网络化、智能化改造，深化数字技术在农业、工业、服务业等领域的应用，促进数字技术与实体经济深度融合，培育融合新业态、新模式。
（四）加强数字基础设施建设。系统布局新型基础设施，加快5G网络、数据中心、工业互联网等新型基础设施建设，构建高速、移动、安全、泛在的数字基础设施体系。
（五）加强数据资源开发利用。加快数据要素市场化，推进公共数据开放共享，促进数据资源整合和开发利用，保障数据安全，构建数据要素有序流通的制度规则。
（六）培育数字经济新业态新模式。大力发展平台经济、共享经济、算法经济等新业态，推动在线教育、远程医疗、数字文化等新模式发展，促进线上线下融合发展。
三、保障措施
（一）加强组织领导和统筹协调。建立健全党委领导、政府主导、企业主体、社会参与的数字经济发展工作机制，统筹推进数字经济发展。
（二）完善数字经济政策体系。优化数字经济发展的财税、金融、投资、贸易等支持政策，完善促进数字经济发展的法律法规和标准体系，构建适应数字经济发展的监管体系。
（三）加强数字经济人才培养。完善数字经济人才培养体系，加大高端数字人才引进力度，加强数字技能培训，建设高素质数字经济人才队伍。
（四）强化数字经济安全保障。完善数字安全保障体系，强化关键信息基础设施安全保护，加强数据安全和个人信息保护，增强网络安全防护能力。
（五）深化数字经济国际合作。秉持开放、合作、共赢理念，积极参与数字经济国际规则制定，推动构建开放、包容、普惠的数字经济国际合作体系。
本意见自发布之日起施行。
        """
    }
    
    # 创建临时目录存放示例文件
    example_dir = os.path.join("temp", "examples")
    os.makedirs(example_dir, exist_ok=True)
    
    file_paths = []
    for filename, content in example_files.items():
        file_path = os.path.join(example_dir, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content.strip())
        file_paths.append(file_path)
    
    return file_paths, example_files

def generate_random_text(n_words=500, policy=True):
    """生成随机文本用于测试"""
    # 政策相关词汇
    policy_words = ["政策", "发展", "战略", "规划", "实施", "推进", "创新", "改革", "促进",
                    "建设", "加强", "完善", "健全", "优化", "提升", "深化", "保障", "支持",
                    "体系", "机制", "制度", "措施", "方案", "行动", "工程", "项目", "计划"]
    
    # 常用词汇
    common_words = ["的", "和", "在", "是", "了", "有", "为", "以", "与", "不", "这", "我们",
                   "要", "对", "从", "由", "到", "上", "中", "下", "各", "种", "等", "多"]
    
    # 生成随机文本
    if policy:
        word_list = policy_words * 3 + common_words
        title = f"随机政策文本_{random.randint(1000, 9999)}"
    else:
        word_list = common_words * 3
        title = f"随机文本_{random.randint(1000, 9999)}"
    
    text = title + "\n\n"
    
    # 生成几个段落
    for i in range(5):
        paragraph = ""
        # 每段句子数
        for j in range(random.randint(3, 6)):
            # 每句词数
            sentence_length = random.randint(10, 20)
            sentence = "".join(random.choices(word_list, k=sentence_length))
            paragraph += sentence + "。"
        text += paragraph + "\n\n"
    
    return text, title

def generate_random_data(n_files=5, show_progress=True):
    """生成随机数据集用于测试（支持大规模生成）"""
    random_dir = os.path.join("temp", "random")
    os.makedirs(random_dir, exist_ok=True)
    
    file_paths = []
    file_contents = {}
    
    # 使用批量处理提高效率
    progress_bar = None
    status_text = None
    if show_progress and n_files > 100:
        progress_bar = st.progress(0)
        status_text = st.empty()
    
    for i in range(n_files):
        # 更新进度
        if progress_bar and i % 50 == 0:
            progress = (i + 1) / n_files
            progress_bar.progress(progress)
            status_text.text(f"生成文件 {i+1}/{n_files}")
        
        text, title = generate_random_text(n_words=random.randint(300, 800))
        filename = f"{title}_{i}.txt"  # 添加索引避免重名
        file_path = os.path.join(random_dir, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        file_paths.append(file_path)
        file_contents[filename] = text
        
        # 每处理500个文件进行一次垃圾回收
        if i > 0 and i % 500 == 0:
            gc.collect()
    
    if progress_bar:
        progress_bar.progress(1.0)
        status_text.text("生成完成")
    
    return file_paths, file_contents

def detect_encoding(file_path):
    """检测文件编码"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'latin-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                f.read()
            return encoding
        except UnicodeDecodeError:
            continue
    return 'utf-8'  # 默认返回utf-8

def read_file_with_encoding(file_path):
    """使用适当的编码读取文件"""
    encoding = detect_encoding(file_path)
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            return f.read()
    except Exception as e:
        log_message(f"读取文件 {os.path.basename(file_path)} 失败: {str(e)}", level="error")
        return ""

def read_docx_file(file_path):
    """读取docx文件内容"""
    if not DOCX_AVAILABLE:
        log_message(f"无法读取 {os.path.basename(file_path)}: python-docx未安装", level="error")
        return ""
    
    try:
        doc = Document(file_path)
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        
        content = '\n'.join(paragraphs)
        return content
    except Exception as e:
        log_message(f"读取docx文件 {os.path.basename(file_path)} 失败: {str(e)}", level="error")
        return ""

def read_docx_from_bytes(file_bytes, filename):
    """从字节流读取docx文件内容"""
    if not DOCX_AVAILABLE:
        log_message(f"无法读取 {filename}: python-docx未安装", level="error")
        return ""
    
    try:
        import io
        doc = Document(io.BytesIO(file_bytes))
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        
        content = '\n'.join(paragraphs)
        return content
    except Exception as e:
        log_message(f"读取docx文件 {filename} 失败: {str(e)}", level="error")
        return ""

def read_pdf_file(file_path):
    """读取PDF文件内容"""
    if PDF_LIBRARY is None:
        log_message(f"无法读取 {os.path.basename(file_path)}: PDF库未安装", level="error")
        return ""
    
    try:
        if PDF_LIBRARY == "pdfplumber":
            import pdfplumber
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            return '\n'.join(text_content)
        
        elif PDF_LIBRARY == "PyPDF2":
            import PyPDF2
            text_content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            return '\n'.join(text_content)
    
    except Exception as e:
        log_message(f"读取PDF文件 {os.path.basename(file_path)} 失败: {str(e)}", level="error")
        return ""

def read_pdf_from_bytes(file_bytes, filename):
    """从字节流读取PDF文件内容"""
    if PDF_LIBRARY is None:
        log_message(f"无法读取 {filename}: PDF库未安装", level="error")
        return ""
    
    try:
        import io
        
        if PDF_LIBRARY == "pdfplumber":
            import pdfplumber
            text_content = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            return '\n'.join(text_content)
        
        elif PDF_LIBRARY == "PyPDF2":
            import PyPDF2
            text_content = []
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            return '\n'.join(text_content)
    
    except Exception as e:
        log_message(f"读取PDF文件 {filename} 失败: {str(e)}", level="error")
        return ""

def get_supported_extensions():
    """获取支持的文件扩展名列表"""
    extensions = ['.txt']
    if DOCX_AVAILABLE:
        extensions.extend(['.docx', '.doc'])
    if PDF_LIBRARY is not None:
        extensions.append('.pdf')
    return extensions

def is_supported_file(filename):
    """检查文件是否为支持的格式"""
    ext = os.path.splitext(filename.lower())[1]
    return ext in get_supported_extensions()

def read_any_document(file_path):
    """根据文件类型读取文档内容"""
    ext = os.path.splitext(file_path.lower())[1]
    
    if ext == '.txt':
        return read_file_with_encoding(file_path)
    elif ext in ['.docx', '.doc']:
        return read_docx_file(file_path)
    elif ext == '.pdf':
        return read_pdf_file(file_path)
    else:
        log_message(f"不支持的文件格式: {ext}", level="warning")
        return ""

def render_data_loader():
    """渲染数据加载模块"""
    st.header("数据加载")
    
    # 功能介绍与操作手册
    with st.expander("📖 功能介绍与操作手册", expanded=False):
        st.markdown("""
        ## 📂 数据加载模块
        
        **功能概述**：导入待分析的政策文本文件，是LDA主题模型分析的第一步。
        
        ---
        
        ### 🎯 使用场景
        
        | 场景 | 推荐方式 | 说明 |
        |------|----------|------|
        | 初次体验系统 | 示例数据 | 快速了解系统功能，无需准备数据 |
        | 正式研究分析 | 上传文件 | 上传自己的政策文本进行分析 |
        | 系统性能测试 | 随机数据 | 测试系统处理大规模数据的能力 |
        | 学术论文研究 | 上传文件 | 上传研究对象的政策文本集 |
        
        ---
        
        ### 📋 操作步骤
        
        **方式一：上传文件**
        1. 选择"上传文件"选项
        2. 点击上传区域，选择TXT文件或ZIP压缩包
        3. 点击"开始加载"按钮
        4. 等待文件处理完成
        
        **方式二：使用示例数据**
        1. 选择"使用示例数据"选项
        2. 点击"加载示例数据"按钮
        3. 系统将加载3份预置的政策文件
        
        **方式三：生成随机数据**
        1. 选择"生成随机数据"选项
        2. 设置要生成的文件数量（或使用快捷选项）
        3. 点击"生成随机数据"按钮
        
        ---
        
        ### 📁 支持的文件格式
        
        | 格式 | 说明 | 编码支持 |
        |------|------|----------|
        | .txt | 纯文本文件 | UTF-8, GBK, GB2312, GB18030, Big5 |
        | .docx | Word文档（推荐） | 自动解析 |
        | .doc | 旧版Word文档 | 自动解析 |
        | .pdf | PDF文档 | 自动提取文本 |
        | .zip | 压缩包（内含多个文档） | 自动检测 |
        
        **注意**：
        - DOC/DOCX支持需要安装python-docx库：`pip install python-docx`
        - PDF支持需要安装pdfplumber或PyPDF2库：`pip install pdfplumber` 或 `pip install PyPDF2`
        - ZIP压缩包可包含TXT、DOC、DOCX、PDF混合文件
        - 扫描版PDF（图片PDF）无法提取文本，建议使用文字版PDF
        
        ---
        
        ### ⚙️ 参数说明
        
        | 参数 | 范围 | 默认值 | 说明 |
        |------|------|--------|------|
        | 最大文件数 | 1-10000 | 10000 | 系统支持的最大文件数量 |
        | 批量处理大小 | - | 100 | 每批处理的文件数量 |
        | 内存警告阈值 | - | 1000 | 超过此数量显示内存警告 |
        
        ---
        
        ### 💡 使用建议
        
        **数据准备建议：**
        - 每个TXT文件包含一份完整的政策文档
        - 文件名建议包含政策名称或编号，便于后续识别
        - 确保文本内容为纯文本格式，避免包含特殊格式
        
        **大规模数据处理：**
        - 超过100个文件建议使用ZIP压缩包上传
        - 超过1000个文件时注意内存使用
        - 处理完成后可点击"清空数据"释放内存
        
        **学术研究建议：**
        - 建议收集同一领域、同一时期的政策文本
        - 文本数量建议在50-500份之间
        - 每份文本长度建议在500-5000字之间
        
        ---
        
        ### ❓ 常见问题
        
        **Q: 文件上传失败怎么办？**
        A: 检查文件编码是否为支持的格式，尝试将文件转换为UTF-8编码。
        
        **Q: 系统显示内存警告怎么办？**
        A: 可以减少加载的文件数量，或在分析完成后点击"清空数据"释放内存。
        
        **Q: 如何判断数据加载是否成功？**
        A: 加载成功后会显示文件列表和统计信息，可以预览文件内容确认。
        """)
    
    # 数据加载选项
    data_source = st.radio(
        "选择数据来源",
        ["上传文件", "使用示例数据", "生成随机数据"],
        horizontal=True,
        key="data_source_radio"
    )
    
    if data_source == "上传文件":
        st.subheader("上传政策文件")
        
        # 构建支持的文件类型列表和提示信息
        supported_formats = ["TXT", "ZIP"]
        
        if DOCX_AVAILABLE:
            supported_formats.extend(["DOC", "DOCX"])
        
        if PDF_LIBRARY is not None:
            supported_formats.append("PDF")
        
        format_str = "、".join(supported_formats)
        st.info(f"💡 提示：支持{format_str}文件，大量文件建议打包成ZIP压缩包上传，支持最多10000个文件")
        
        # 不限制文件类型，让用户上传任何文件，在处理时检查
        uploaded_files = st.file_uploader(
            "上传文本文件或ZIP压缩包", 
            accept_multiple_files=True,
            key="data_files_uploader",
            help="支持TXT、DOC、DOCX、PDF、ZIP格式"
        )
        
        if uploaded_files and st.button("开始加载", key="load_uploaded_files"):
            with st.spinner("正在处理上传文件..."):
                st.session_state.file_contents = {}
                st.session_state.file_names = []
                st.session_state.raw_texts = []
                
                total_files = len(uploaded_files)
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                processed_count = 0
                skipped_count = 0
                
                for i, uploaded_file in enumerate(uploaded_files):
                    status_text.text(f"处理文件 {i+1}/{total_files}: {uploaded_file.name}")
                    
                    if uploaded_file.name.endswith('.zip'):
                        # 处理ZIP文件
                        with tempfile.TemporaryDirectory() as temp_dir:
                            zip_path = os.path.join(temp_dir, uploaded_file.name)
                            with open(zip_path, 'wb') as f:
                                f.write(uploaded_file.getbuffer())
                            
                            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                                # 获取ZIP中支持的文件列表
                                supported_files = [f for f in zip_ref.namelist() if is_supported_file(f)]
                                total_supported = len(supported_files)
                                
                                zip_ref.extractall(temp_dir)
                                
                                # 批量处理ZIP中的文件
                                for j, file in enumerate(supported_files):
                                    if j % 50 == 0:
                                        progress = (i + j/max(total_supported, 1)) / total_files
                                        progress_bar.progress(min(progress, 1.0))
                                        status_text.text(f"处理ZIP文件 {i+1}/{total_files}，内部文件 {j+1}/{total_supported}")
                                    
                                    file_path = os.path.join(temp_dir, file)
                                    if os.path.isfile(file_path):
                                        content = read_any_document(file_path)
                                        if content and content.strip():
                                            filename = os.path.basename(file)
                                            st.session_state.file_contents[filename] = content
                                            st.session_state.file_names.append(filename)
                                            st.session_state.raw_texts.append(content)
                                            processed_count += 1
                                        else:
                                            skipped_count += 1
                                    
                                    # 定期垃圾回收
                                    if j > 0 and j % 500 == 0:
                                        gc.collect()
                    
                    elif uploaded_file.name.lower().endswith('.txt'):
                        # 处理单个TXT文件
                        content = uploaded_file.read().decode('utf-8', errors='replace')
                        if content and content.strip():
                            st.session_state.file_contents[uploaded_file.name] = content
                            st.session_state.file_names.append(uploaded_file.name)
                            st.session_state.raw_texts.append(content)
                            processed_count += 1
                        else:
                            skipped_count += 1
                    
                    elif uploaded_file.name.lower().endswith(('.docx', '.doc')):
                        # 处理DOC/DOCX文件
                        if DOCX_AVAILABLE:
                            file_bytes = uploaded_file.read()
                            content = read_docx_from_bytes(file_bytes, uploaded_file.name)
                            if content and content.strip():
                                st.session_state.file_contents[uploaded_file.name] = content
                                st.session_state.file_names.append(uploaded_file.name)
                                st.session_state.raw_texts.append(content)
                                processed_count += 1
                            else:
                                skipped_count += 1
                                log_message(f"文件 {uploaded_file.name} 内容为空，已跳过", level="warning")
                        else:
                            skipped_count += 1
                            log_message(f"无法处理 {uploaded_file.name}: python-docx未安装", level="warning")
                    
                    elif uploaded_file.name.lower().endswith('.pdf'):
                        # 处理PDF文件
                        if PDF_LIBRARY is not None:
                            file_bytes = uploaded_file.read()
                            content = read_pdf_from_bytes(file_bytes, uploaded_file.name)
                            if content and content.strip():
                                st.session_state.file_contents[uploaded_file.name] = content
                                st.session_state.file_names.append(uploaded_file.name)
                                st.session_state.raw_texts.append(content)
                                processed_count += 1
                            else:
                                skipped_count += 1
                                log_message(f"文件 {uploaded_file.name} 内容为空或无法提取文本，已跳过", level="warning")
                        else:
                            skipped_count += 1
                            log_message(f"无法处理 {uploaded_file.name}: PDF库未安装", level="warning")
                    
                    progress_bar.progress((i + 1) / total_files)
                
                progress_bar.progress(1.0)
                status_text.text("文件加载完成")
                update_progress(1.0, "文件加载完成")
                log_message(f"已加载 {processed_count} 个文件", level="success")
                
                # 显示加载结果
                if skipped_count > 0:
                    st.success(f"成功加载 {processed_count} 个文件，跳过 {skipped_count} 个空文件或不支持的文件")
                else:
                    st.success(f"成功加载 {processed_count} 个文件")
                
                # 内存警告
                if processed_count > MEMORY_WARNING_THRESHOLD:
                    st.warning(f"⚠️ 已加载 {processed_count} 个文件，请注意内存使用。建议在处理完成后清理缓存。")
                
                # 清空上传的文件
                st.session_state.uploaded_files = None
    
    elif data_source == "使用示例数据":
        st.subheader("示例政策数据")
        st.info("这些是预设的政策文件示例，用于系统功能演示。")
        
        if st.button("加载示例数据", key="load_example_data"):
            with st.spinner("正在加载示例数据..."):
                file_paths, file_contents = load_example_data()
                
                st.session_state.file_contents = file_contents
                st.session_state.file_names = list(file_contents.keys())
                st.session_state.raw_texts = list(file_contents.values())
                
                log_message(f"已加载 {len(file_paths)} 个示例文件", level="success")
                st.success(f"成功加载 {len(file_paths)} 个示例文件")
    
    elif data_source == "生成随机数据":
        st.subheader("生成随机测试数据")
        
        # 扩展滑块范围以支持大规模测试
        col1, col2 = st.columns(2)
        with col1:
            n_files = st.number_input(
                "生成文件数量", 
                min_value=3, 
                max_value=MAX_FILES, 
                value=100,
                step=100,
                help=f"支持生成3-{MAX_FILES}个测试文件",
                key="n_files_input"
            )
        with col2:
            # 快捷选项
            quick_options = st.selectbox(
                "快捷选项",
                ["自定义", "100个(快速测试)", "500个(中等规模)", "2000个(大规模)", "5000个(压力测试)"],
                key="quick_file_options"
            )
            if quick_options != "自定义":
                n_files = int(quick_options.split("个")[0])
        
        # 大规模生成警告
        if n_files > 1000:
            st.warning(f"⚠️ 即将生成 {n_files} 个文件，这可能需要较长时间和较多内存。")
        
        if st.button("生成随机数据", key="generate_random_data"):
            start_time = time.time()
            
            with st.spinner(f"正在生成 {n_files} 个随机文件..."):
                file_paths, file_contents = generate_random_data(n_files, show_progress=True)
                
                st.session_state.file_contents = file_contents
                st.session_state.file_names = list(file_contents.keys())
                st.session_state.raw_texts = list(file_contents.values())
                
                elapsed = time.time() - start_time
                log_message(f"已生成 {len(file_paths)} 个随机文件，耗时 {elapsed:.1f}秒", level="success")
                st.success(f"成功生成 {len(file_paths)} 个随机文件，耗时 {elapsed:.1f}秒")
                
                # 垃圾回收
                gc.collect()
    
    # 显示已加载的文件
    if st.session_state.get("file_contents"):
        st.subheader("已加载的文件")
        
        # 添加清空数据按钮
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            file_count = len(st.session_state.file_names)
            total_chars = sum(len(text) for text in st.session_state.raw_texts)
            st.metric("文件统计", f"{file_count} 个文件，共 {total_chars:,} 字符")
        with col3:
            if st.button("🗑️ 清空数据", key="clear_loaded_data", type="secondary"):
                st.session_state.file_contents = {}
                st.session_state.file_names = []
                st.session_state.raw_texts = []
                st.session_state.texts = None
                st.session_state.dictionary = None
                st.session_state.corpus = None
                st.session_state.lda_model = None
                st.session_state.training_complete = False
                st.session_state.pyldavis_html = None
                st.session_state.wordcloud_images = {}
                gc.collect()  # 强制垃圾回收
                log_message("已清空所有加载的数据", level="warning")
                st.rerun()
        
        # 大量文件时只显示摘要
        file_count = len(st.session_state.file_names)
        if file_count > 100:
            st.info(f"📊 已加载 {file_count} 个文件（数据量较大，仅显示前100个）")
            display_names = st.session_state.file_names[:100]
            display_texts = st.session_state.raw_texts[:100]
        else:
            display_names = st.session_state.file_names
            display_texts = st.session_state.raw_texts
        
        df = pd.DataFrame({
            "文件名": display_names,
            "文本长度": [len(text) for text in display_texts]
        })
        
        st.dataframe(df, use_container_width=True, height=300)
        
        # 文件预览
        if st.session_state.file_names:
            with st.expander("文件内容预览", expanded=False):
                # 限制预览选项数量
                preview_options = st.session_state.file_names[:100] if file_count > 100 else st.session_state.file_names
                preview_file = st.selectbox("选择文件预览", preview_options, key="data_preview_file_select")
                if preview_file:
                    content = st.session_state.file_contents[preview_file]
                    # 限制预览内容长度
                    max_preview_length = 1000
                    if len(content) > max_preview_length:
                        preview_content = content[:max_preview_length] + "..."
                    else:
                        preview_content = content
                    
                    st.text_area("文件内容", preview_content, height=300) 