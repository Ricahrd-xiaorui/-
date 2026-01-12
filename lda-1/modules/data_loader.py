import streamlit as st
import os
import zipfile
import tempfile
import random
import string
import time
import re
import pandas as pd
import gc
from pathlib import Path
from utils.session_state import get_session_state, log_message, update_progress

# å°è¯•å¯¼å…¥python-docxï¼Œå¦‚æžœå¤±è´¥åˆ™è®¾ç½®æ ‡å¿—
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    log_message("python-docxæœªå®‰è£…ï¼Œæ— æ³•å¤„ç†.docxæ–‡ä»¶", level="warning")

# å¤§è§„æ¨¡æ–‡ä»¶å¤„ç†é…ç½®
MAX_FILES = 10000  # æœ€å¤§æ”¯æŒæ–‡ä»¶æ•°
BATCH_SIZE = 100   # æ‰¹é‡å¤„ç†å¤§å°
MEMORY_WARNING_THRESHOLD = 1000  # è¶…è¿‡æ­¤æ•°é‡æ˜¾ç¤ºå†…å­˜è­¦å‘Š

def load_example_data():
    """åŠ è½½ç¤ºä¾‹æ”¿ç­–æ•°æ®"""
    example_files = {
        "æ”¿ç­–1_ç§‘æŠ€åˆ›æ–°æ”¿ç­–.txt": """
ä¸ºå…¨é¢è´¯å½»å…šçš„åä¹å¤§å’Œåä¹å±ŠäºŒä¸­ã€ä¸‰ä¸­ã€å››ä¸­ã€äº”ä¸­å…¨ä¼šç²¾ç¥žï¼Œè½å®žã€Šå›½å®¶åˆ›æ–°é©±åŠ¨å‘å±•æˆ˜ç•¥çº²è¦ã€‹ï¼ŒæŽ¨åŠ¨ç§‘æŠ€åˆ›æ–°ä¸Žäº§ä¸šå‘å±•æ·±åº¦èžåˆï¼ŒåŠ å¿«å»ºè®¾ç§‘æŠ€å¼ºå›½ï¼ŒçŽ°æå‡ºå¦‚ä¸‹æ„è§ã€‚
ä¸€ã€æ€»ä½“è¦æ±‚
ï¼ˆä¸€ï¼‰æŒ‡å¯¼æ€æƒ³ã€‚ä»¥ä¹ è¿‘å¹³æ–°æ—¶ä»£ä¸­å›½ç‰¹è‰²ç¤¾ä¼šä¸»ä¹‰æ€æƒ³ä¸ºæŒ‡å¯¼ï¼Œæ·±å…¥è´¯å½»å…šçš„åä¹å¤§å’Œåä¹å±ŠäºŒä¸­ã€ä¸‰ä¸­ã€å››ä¸­ã€äº”ä¸­å…¨ä¼šç²¾ç¥žï¼ŒåšæŒåˆ›æ–°åœ¨æˆ‘å›½çŽ°ä»£åŒ–å»ºè®¾å…¨å±€ä¸­çš„æ ¸å¿ƒåœ°ä½ï¼ŒæŠŠç§‘æŠ€è‡ªç«‹è‡ªå¼ºä½œä¸ºå›½å®¶å‘å±•çš„æˆ˜ç•¥æ”¯æ’‘ï¼Œæ·±å…¥å®žæ–½åˆ›æ–°é©±åŠ¨å‘å±•æˆ˜ç•¥ï¼Œå®Œå–„å›½å®¶åˆ›æ–°ä½“ç³»ï¼ŒåŠ å¿«å»ºè®¾ç§‘æŠ€å¼ºå›½ï¼Œå®žçŽ°é«˜æ°´å¹³ç§‘æŠ€è‡ªç«‹è‡ªå¼ºã€‚
ï¼ˆäºŒï¼‰ä¸»è¦ç›®æ ‡ã€‚åˆ°2025å¹´ï¼Œå»ºæˆä¸€æ‰¹å„å…·ç‰¹è‰²ã€ä¼˜åŠ¿äº’è¡¥ã€ç»“æž„åˆç†çš„ç§‘æŠ€åˆ›æ–°åŸºåœ°ï¼Œå½¢æˆåŒºåŸŸåˆ›æ–°é«˜åœ°å’Œç‰¹è‰²åˆ›æ–°é›†ç¾¤ï¼Œæå‡åŒºåŸŸç§‘æŠ€åˆ›æ–°èƒ½åŠ›ã€‚åŠ å¿«æŽ¨è¿›ç§‘æŠ€æˆæžœè½¬åŒ–åº”ç”¨ï¼Œä¿ƒè¿›äº§ä¸šé“¾åˆ›æ–°é“¾æ·±åº¦èžåˆï¼Œæž„å»ºçŽ°ä»£äº§ä¸šæŠ€æœ¯ä½“ç³»ï¼Œå¤§å¹…æå‡ä¼ä¸šåˆ›æ–°èƒ½åŠ›ï¼Œæ”¯æ’‘å¼•é¢†é«˜è´¨é‡å‘å±•ã€‚
äºŒã€é‡ç‚¹ä»»åŠ¡
ï¼ˆä¸€ï¼‰åŠ å¼ºåŽŸåˆ›æ€§å¼•é¢†æ€§ç§‘æŠ€æ”»å…³ã€‚çž„å‡†ä¸–ç•Œç§‘æŠ€å‰æ²¿ã€ç»æµŽä¸»æˆ˜åœºã€å›½å®¶é‡å¤§éœ€æ±‚ã€äººæ°‘ç”Ÿå‘½å¥åº·ï¼ŒåŠ å¼ºåŸºç¡€ç ”ç©¶ã€æ³¨é‡åŽŸå§‹åˆ›æ–°ï¼Œä¼˜åŒ–å­¦ç§‘å¸ƒå±€å’Œç ”å‘å¸ƒå±€ï¼ŒåŠ å¼ºåŸºç¡€å‰æ²¿äº¤å‰ç ”ç©¶ï¼Œä¿ƒè¿›å­¦ç§‘äº¤å‰èžåˆã€‚
ï¼ˆäºŒï¼‰ä¿ƒè¿›ç§‘æŠ€åˆ›æ–°ä¸Žäº§ä¸šå‘å±•æ·±åº¦èžåˆã€‚å®Œå–„äº§å­¦ç ”ååŒåˆ›æ–°ä½“ç³»ï¼Œå¤§åŠ›æŽ¨è¿›ç§‘æŠ€æˆæžœè½¬åŒ–åº”ç”¨ï¼Œæå‡äº§ä¸šé“¾çŽ°ä»£åŒ–æ°´å¹³ã€‚åŠ å¿«æ•°å­—åŒ–æ™ºèƒ½åŒ–ç»¿è‰²åŒ–è½¬åž‹ï¼ŒåŠ å¼ºå…±æ€§æŠ€æœ¯å¹³å°å»ºè®¾ï¼Œä¿ƒè¿›ä¸­å°ä¼ä¸šåˆ›æ–°å‘å±•ï¼ŒæŽ¨åŠ¨é‡ç‚¹äº§ä¸šé›†ç¾¤å‘å±•ã€‚
ï¼ˆä¸‰ï¼‰åŸ¹å…»é€ å°±é«˜æ°´å¹³ç§‘æŠ€äººæ‰é˜Ÿä¼ã€‚å¼ºåŒ–æˆ˜ç•¥ç§‘æŠ€äººæ‰åŠ›é‡ï¼ŒåŸ¹å…»å…·æœ‰å›½é™…æ°´å¹³çš„æˆ˜ç•¥ç§‘æŠ€äººæ‰ã€ç§‘æŠ€é¢†å†›äººæ‰ã€é’å¹´ç§‘æŠ€äººæ‰å’Œé«˜æ°´å¹³åˆ›æ–°å›¢é˜Ÿã€‚å®Œå–„äººæ‰å‘çŽ°ã€åŸ¹å…»ã€æ¿€åŠ±æœºåˆ¶ï¼Œåˆ›æ–°äººæ‰è¯„ä»·æœºåˆ¶ï¼Œæž„å»ºå…·æœ‰å›½é™…ç«žäº‰åŠ›çš„å¼•æ‰ç”¨æ‰åˆ¶åº¦ä½“ç³»ã€‚
ä¸‰ã€ä¿éšœæŽªæ–½
ï¼ˆä¸€ï¼‰åŠ å¤§ç§‘æŠ€åˆ›æ–°æŠ•å…¥ã€‚å®Œå–„æ”¿åºœæŠ•å…¥ä¸ºä¸»ã€ç¤¾ä¼šå¤šæ¸ é“æŠ•å…¥æœºåˆ¶ï¼Œä¿ƒè¿›ç§‘æŠ€æŠ•å…¥æŒç»­ç¨³å®šå¢žé•¿ã€‚å¼•å¯¼ä¼ä¸šåŠ å¤§ç ”å‘æŠ•å…¥ï¼Œé¼“åŠ±ç¤¾ä¼šåŠ›é‡æŠ•å…¥åŸºç¡€ç ”ç©¶ï¼Œå»ºç«‹å¥å…¨åŸºç¡€ç ”ç©¶å¤šå…ƒæŠ•å…¥æœºåˆ¶ã€‚
ï¼ˆäºŒï¼‰å®Œå–„ç§‘æŠ€åˆ›æ–°ä½“åˆ¶æœºåˆ¶ã€‚æ·±åŒ–ç§‘æŠ€ä½“åˆ¶æ”¹é©ï¼Œå¥å…¨ç¤¾ä¼šä¸»ä¹‰å¸‚åœºç»æµŽæ¡ä»¶ä¸‹æ–°åž‹ä¸¾å›½ä½“åˆ¶ï¼Œå……åˆ†å‘æŒ¥æ”¿åºœä½œç”¨å’Œå¸‚åœºæœºåˆ¶ä½œç”¨ï¼Œå®Œå–„ç§‘æŠ€æ²»ç†ä½“ç³»ï¼Œä¼˜åŒ–ç§‘æŠ€èµ„æºé…ç½®ï¼Œä¿ƒè¿›ç§‘æŠ€äº‹ä¸šå…¨é¢å‘å±•ã€‚
ï¼ˆä¸‰ï¼‰å¼ºåŒ–çŸ¥è¯†äº§æƒä¿æŠ¤å’Œè¿ç”¨ã€‚å¥å…¨çŸ¥è¯†äº§æƒä¿æŠ¤æ³•å¾‹æ³•è§„ï¼Œå®Œå–„çŸ¥è¯†äº§æƒä¿æŠ¤ä½“ç³»ï¼ŒåŠ å¼ºçŸ¥è¯†äº§æƒä¿æŠ¤æ‰§æ³•ã€‚å¥å…¨çŸ¥è¯†äº§æƒè¿è¥æœåŠ¡ä½“ç³»ï¼ŒåŸ¹è‚²çŸ¥è¯†äº§æƒå¯†é›†åž‹äº§ä¸šï¼Œæå‡çŸ¥è¯†äº§æƒå…¬å…±æœåŠ¡æ°´å¹³ã€‚
ï¼ˆå››ï¼‰æ·±åŒ–ç§‘æŠ€å¼€æ”¾åˆä½œã€‚ä¸»åŠ¨èžå…¥å…¨çƒåˆ›æ–°ç½‘ç»œï¼ŒåšæŒ"å¼•è¿›æ¥"å’Œ"èµ°å‡ºåŽ»"ç›¸ç»“åˆï¼Œä¿ƒè¿›å›½é™…ç§‘æŠ€äº¤æµåˆä½œï¼Œæé«˜ç§‘æŠ€å¼€æ”¾åˆä½œæ°´å¹³ã€‚ä¼˜åŒ–åˆ›æ–°åˆ›ä¸šçŽ¯å¢ƒï¼ŒæŽ¨åŠ¨å½¢æˆå¼€æ”¾ã€èžåˆã€å…±äº«çš„åˆ›æ–°ç”Ÿæ€ã€‚
æœ¬æ„è§è‡ªå‘å¸ƒä¹‹æ—¥èµ·æ–½è¡Œã€‚
        """,
        "æ”¿ç­–2_ä¹¡æ‘æŒ¯å…´æˆ˜ç•¥.txt": """
ä¸ºæ·±å…¥è´¯å½»å…šçš„åä¹å¤§å’Œåä¹å±ŠäºŒä¸­ã€ä¸‰ä¸­ã€å››ä¸­ã€äº”ä¸­å…¨ä¼šç²¾ç¥žï¼Œè½å®žã€Šä¸­å…±ä¸­å¤®å›½åŠ¡é™¢å…³äºŽå®žæ–½ä¹¡æ‘æŒ¯å…´æˆ˜ç•¥çš„æ„è§ã€‹ï¼Œå…¨é¢æŽ¨è¿›ä¹¡æ‘æŒ¯å…´æˆ˜ç•¥å®žæ–½ï¼ŒåŠ å¿«å†œä¸šå†œæ‘çŽ°ä»£åŒ–ï¼ŒçŽ°æå‡ºå¦‚ä¸‹æ„è§ã€‚
ä¸€ã€æ€»ä½“è¦æ±‚
ï¼ˆä¸€ï¼‰æŒ‡å¯¼æ€æƒ³ã€‚ä»¥ä¹ è¿‘å¹³æ–°æ—¶ä»£ä¸­å›½ç‰¹è‰²ç¤¾ä¼šä¸»ä¹‰æ€æƒ³ä¸ºæŒ‡å¯¼ï¼Œæ·±å…¥è´¯å½»å…šçš„åä¹å¤§å’Œåä¹å±ŠäºŒä¸­ã€ä¸‰ä¸­ã€å››ä¸­ã€äº”ä¸­å…¨ä¼šç²¾ç¥žï¼ŒåšæŒå†œä¸šå†œæ‘ä¼˜å…ˆå‘å±•ï¼ŒæŒ‰ç…§äº§ä¸šå…´æ—ºã€ç”Ÿæ€å®œå±…ã€ä¹¡é£Žæ–‡æ˜Žã€æ²»ç†æœ‰æ•ˆã€ç”Ÿæ´»å¯Œè£•çš„æ€»è¦æ±‚ï¼Œå»ºç«‹å¥å…¨åŸŽä¹¡èžåˆå‘å±•ä½“åˆ¶æœºåˆ¶å’Œæ”¿ç­–ä½“ç³»ï¼ŒåŠ å¿«æŽ¨è¿›å†œä¸šå†œæ‘çŽ°ä»£åŒ–ã€‚
ï¼ˆäºŒï¼‰ä¸»è¦ç›®æ ‡ã€‚åˆ°2025å¹´ï¼Œä¹¡æ‘æŒ¯å…´å–å¾—é‡è¦è¿›å±•ï¼Œå†œä¸šåŸºç¡€æ›´åŠ ç¨³å›ºï¼Œå†œæ‘åŸºç¡€è®¾æ–½å’Œå…¬å…±æœåŠ¡ä½“ç³»åŸºæœ¬å¥å…¨ï¼ŒåŸŽä¹¡èžåˆå‘å±•ä½“åˆ¶æœºåˆ¶åŸºæœ¬å»ºç«‹ã€‚åˆ°2035å¹´ï¼Œä¹¡æ‘æŒ¯å…´å–å¾—å†³å®šæ€§è¿›å±•ï¼Œå†œä¸šå†œæ‘çŽ°ä»£åŒ–åŸºæœ¬å®žçŽ°ã€‚åˆ°2050å¹´ï¼Œä¹¡æ‘å…¨é¢æŒ¯å…´ï¼Œå†œä¸šå¼ºã€å†œæ‘ç¾Žã€å†œæ°‘å¯Œå…¨é¢å®žçŽ°ã€‚
äºŒã€é‡ç‚¹ä»»åŠ¡
ï¼ˆä¸€ï¼‰å®žæ–½ä¹¡æ‘äº§ä¸šæŒ¯å…´è¡ŒåŠ¨ã€‚å¤§åŠ›å‘å±•ä¹¡æ‘äº§ä¸šï¼Œåšå¼ºåšä¼˜å†œä¸šï¼ŒåŸ¹è‚²æ–°äº§ä¸šæ–°ä¸šæ€ï¼Œæ‹“å±•å†œä¸šå¤šç§åŠŸèƒ½ï¼ŒæŽ¨åŠ¨äº§ä¸šèžåˆå‘å±•ï¼Œæž„å»ºçŽ°ä»£ä¹¡æ‘äº§ä¸šä½“ç³»ï¼ŒæŽ¨åŠ¨å†œæ‘ä¸€äºŒä¸‰äº§ä¸šèžåˆå‘å±•ã€‚
ï¼ˆäºŒï¼‰å®žæ–½ä¹¡æ‘äººæ‰æŒ¯å…´è¡ŒåŠ¨ã€‚åŠ å¼ºå†œæ‘äººæ‰é˜Ÿä¼å»ºè®¾ï¼ŒåŸ¹è‚²æ–°åž‹èŒä¸šå†œæ°‘ï¼ŒåŠ å¼ºå†œæ‘ä¸“ä¸šäººæ‰é˜Ÿä¼å»ºè®¾ï¼Œå‘æŒ¥ç§‘æŠ€äººæ‰æ”¯æ’‘ä½œç”¨ï¼Œé¼“åŠ±ç¤¾ä¼šå„ç•Œäººæ‰æŠ•èº«ä¹¡æ‘å»ºè®¾ï¼Œåˆ›æ–°ä¹¡æ‘äººæ‰åŸ¹è‚²å¼•è¿›ä½¿ç”¨æœºåˆ¶ã€‚
ï¼ˆä¸‰ï¼‰å®žæ–½ä¹¡æ‘æ–‡åŒ–æŒ¯å…´è¡ŒåŠ¨ã€‚åŠ å¼ºå†œæ‘æ€æƒ³é“å¾·å»ºè®¾ï¼Œå¼˜æ‰¬å’Œè·µè¡Œç¤¾ä¼šä¸»ä¹‰æ ¸å¿ƒä»·å€¼è§‚ï¼Œä¼ æ‰¿å‘å±•ä¼˜ç§€ä¼ ç»Ÿæ–‡åŒ–ï¼ŒåŠ å¼ºå†œæ‘å…¬å…±æ–‡åŒ–å»ºè®¾ï¼Œå¼€å±•ç§»é£Žæ˜“ä¿—è¡ŒåŠ¨ï¼ŒæŽ¨åŠ¨å½¢æˆæ–‡æ˜Žä¹¡é£Žã€è‰¯å¥½å®¶é£Žã€æ·³æœ´æ°‘é£Žã€‚
ï¼ˆå››ï¼‰å®žæ–½ä¹¡æ‘ç”Ÿæ€æŒ¯å…´è¡ŒåŠ¨ã€‚åŠ å¼ºå†œæ‘ç”Ÿæ€çŽ¯å¢ƒä¿æŠ¤å’Œæ²»ç†ï¼ŒæŽ¨åŠ¨å†œä¸šç»¿è‰²å‘å±•ï¼Œå¼€å±•å†œæ‘äººå±…çŽ¯å¢ƒæ•´æ²»è¡ŒåŠ¨ï¼ŒåŠ å¼ºå†œæ‘ç”Ÿæ€ä¿æŠ¤ä¸Žä¿®å¤ï¼Œå¥å…¨å†œæ‘ç”Ÿæ€æ–‡æ˜Žåˆ¶åº¦ä½“ç³»ï¼Œå»ºè®¾ç”Ÿæ€å®œå±…ç¾Žä¸½ä¹¡æ‘ã€‚
ï¼ˆäº”ï¼‰å®žæ–½ä¹¡æ‘ç»„ç»‡æŒ¯å…´è¡ŒåŠ¨ã€‚åŠ å¼ºå†œæ‘åŸºå±‚å…šç»„ç»‡å»ºè®¾ï¼Œå¥å…¨ä¹¡æ‘ç»„ç»‡ä½“ç³»ï¼Œå®Œå–„æ‘æ°‘è‡ªæ²»æœºåˆ¶ï¼Œå¥å…¨æ‘çº§è®®äº‹åå•†åˆ¶åº¦ï¼Œå‘æŒ¥å†œæ‘åŸºå±‚å…šç»„ç»‡æˆ˜æ–—å ¡åž’ä½œç”¨ï¼Œå¢žå¼ºä¹¡æ‘æ²»ç†èƒ½åŠ›ã€‚
ï¼ˆå…­ï¼‰å®žæ–½ä¹¡æ‘ç”Ÿæ´»å¯Œè£•è¡ŒåŠ¨ã€‚æŒç»­å¢žåŠ å†œæ°‘æ”¶å…¥ï¼Œä¼˜å…ˆå‘å±•å†œæ‘æ•™è‚²äº‹ä¸šï¼ŒæŽ¨è¿›å¥åº·ä¹¡æ‘å»ºè®¾ï¼Œå¥å…¨å†œæ‘ç¤¾ä¼šä¿éšœä½“ç³»ï¼Œæ”¹å–„å†œæ‘äººå±…çŽ¯å¢ƒï¼Œæé«˜å†œæ‘ç”Ÿæ´»è´¨é‡ï¼Œä¿ƒè¿›å†œæ°‘ç”Ÿæ´»å¯Œè£•ã€‚
ä¸‰ã€ä¿éšœæŽªæ–½
ï¼ˆä¸€ï¼‰åŠ å¼ºå…šå¯¹ä¹¡æ‘æŒ¯å…´çš„é¢†å¯¼ã€‚åšæŒå…šç®¡å†œæ‘å·¥ä½œï¼Œå¥å…¨å…šå§”ç»Ÿä¸€é¢†å¯¼ã€æ”¿åºœè´Ÿè´£ã€å…šå§”å†œæ‘å·¥ä½œéƒ¨é—¨ç»Ÿç­¹åè°ƒçš„å†œæ‘å·¥ä½œé¢†å¯¼ä½“åˆ¶ï¼Œå»ºç«‹å¥å…¨å®žæ–½ä¹¡æ‘æŒ¯å…´æˆ˜ç•¥é¢†å¯¼è´£ä»»åˆ¶ã€‚
ï¼ˆäºŒï¼‰å®Œå–„ä¹¡æ‘æŒ¯å…´æ”¿ç­–ä½“ç³»ã€‚å¥å…¨åŸŽä¹¡èžåˆå‘å±•ä½“åˆ¶æœºåˆ¶ï¼Œå®Œå–„å†œä¸šå†œæ‘æ”¯æŒä¿æŠ¤åˆ¶åº¦ï¼ŒåŠ å¤§ä¹¡æ‘æŒ¯å…´æŠ•å…¥åŠ›åº¦ï¼Œå¼ºåŒ–ä¹¡æ‘æŒ¯å…´äººæ‰æ”¯æ’‘ã€‚
ï¼ˆä¸‰ï¼‰å¼ºåŒ–ä¹¡æ‘æŒ¯å…´æ³•æ²»ä¿éšœã€‚å®Œå–„ä¹¡æ‘æŒ¯å…´æ³•å¾‹æ³•è§„ä½“ç³»ï¼Œå¢žå¼ºä¹¡æ‘æ³•æ²»è§‚å¿µï¼ŒåŠ å¼ºå†œæ‘æ³•æ²»å®£ä¼ æ•™è‚²ï¼Œä¿éšœå†œæ°‘åˆæ³•æƒç›Šï¼ŒæŽ¨è¿›æ³•æ²»ä¹¡æ‘å»ºè®¾ã€‚
ï¼ˆå››ï¼‰åŠ å¼ºä¹¡æ‘æŒ¯å…´è€ƒæ ¸è¯„ä»·ã€‚å»ºç«‹å¥å…¨ä¹¡æ‘æŒ¯å…´ç›‘æµ‹è¯„ä»·æœºåˆ¶ï¼Œå¼€å±•ä¹¡æ‘æŒ¯å…´æˆ˜ç•¥å®žæ–½ç›‘æµ‹è¯„ä»·ï¼Œå°†ä¹¡æ‘æŒ¯å…´æˆ˜ç•¥å®žæ–½æˆæ•ˆçº³å…¥å…šæ”¿é¢†å¯¼ç­å­å’Œé¢†å¯¼å¹²éƒ¨å®žç»©è€ƒæ ¸å†…å®¹ã€‚
æœ¬æ„è§è‡ªå‘å¸ƒä¹‹æ—¥èµ·æ–½è¡Œã€‚
        """,
        "æ”¿ç­–3_æ•°å­—ç»æµŽå‘å±•æ”¿ç­–.txt": """
ä¸ºæ·±å…¥è´¯å½»å…šçš„åä¹å¤§å’Œåä¹å±ŠäºŒä¸­ã€ä¸‰ä¸­ã€å››ä¸­ã€äº”ä¸­å…¨ä¼šç²¾ç¥žï¼Œè½å®žã€Šå›½å®¶æ•°å­—ç»æµŽå‘å±•æˆ˜ç•¥ã€‹ï¼ŒåŠ å¿«æ•°å­—ç»æµŽå‘å±•ï¼ŒæŽ¨åŠ¨æ•°å­—æŠ€æœ¯ä¸Žå®žä½“ç»æµŽæ·±åº¦èžåˆï¼ŒçŽ°æå‡ºå¦‚ä¸‹æ„è§ã€‚
ä¸€ã€æ€»ä½“è¦æ±‚
ï¼ˆä¸€ï¼‰æŒ‡å¯¼æ€æƒ³ã€‚ä»¥ä¹ è¿‘å¹³æ–°æ—¶ä»£ä¸­å›½ç‰¹è‰²ç¤¾ä¼šä¸»ä¹‰æ€æƒ³ä¸ºæŒ‡å¯¼ï¼Œæ·±å…¥è´¯å½»å…šçš„åä¹å¤§å’Œåä¹å±ŠäºŒä¸­ã€ä¸‰ä¸­ã€å››ä¸­ã€äº”ä¸­å…¨ä¼šç²¾ç¥žï¼Œç«‹è¶³æ–°å‘å±•é˜¶æ®µï¼Œè´¯å½»æ–°å‘å±•ç†å¿µï¼Œæž„å»ºæ–°å‘å±•æ ¼å±€ï¼ŒæŽ¨åŠ¨æ•°å­—ç»æµŽå’Œå®žä½“ç»æµŽæ·±åº¦èžåˆï¼Œèµ‹èƒ½ä¼ ç»Ÿäº§ä¸šè½¬åž‹å‡çº§ï¼ŒåŸ¹è‚²æ–°äº§ä¸šæ–°ä¸šæ€æ–°æ¨¡å¼ï¼Œæž„å»ºæ•°å­—ç»æµŽæ–°ä¼˜åŠ¿ã€‚
ï¼ˆäºŒï¼‰ä¸»è¦ç›®æ ‡ã€‚åˆ°2025å¹´ï¼Œæ•°å­—ç»æµŽè¿ˆå‘å…¨é¢æ‰©å±•æœŸï¼Œæ•°å­—ç»æµŽæ ¸å¿ƒäº§ä¸šå¢žåŠ å€¼å GDPæ¯”é‡è¾¾åˆ°10%ï¼Œæ•°å­—åŒ–åˆ›æ–°å¼•é¢†å‘å±•èƒ½åŠ›å¤§å¹…æå‡ï¼Œæ™ºèƒ½åŒ–æ°´å¹³æ˜Žæ˜¾å¢žå¼ºï¼Œæ•°å­—æŠ€æœ¯ä¸Žå®žä½“ç»æµŽæ·±åº¦èžåˆå–å¾—æ˜¾è‘—æˆæ•ˆã€‚åˆ°2035å¹´ï¼Œæ•°å­—ç»æµŽæ•´ä½“ç«žäº‰åŠ›ä½å±…å…¨çƒå‰åˆ—ï¼Œæ•°å­—åŒ–ã€ç½‘ç»œåŒ–ã€æ™ºèƒ½åŒ–å…¨é¢å‘å±•ã€‚
äºŒã€é‡ç‚¹ä»»åŠ¡
ï¼ˆä¸€ï¼‰åŠ å¿«æ•°å­—æŠ€æœ¯åˆ›æ–°çªç ´ã€‚å¼ºåŒ–åŸºç¡€ç ”ç©¶å’ŒåŽŸå§‹åˆ›æ–°ï¼ŒåŠ å¿«äººå·¥æ™ºèƒ½ã€é‡å­ä¿¡æ¯ã€é›†æˆç”µè·¯ã€åŒºå—é“¾ã€6Gã€ç‰©è”ç½‘ç­‰å…³é”®æ•°å­—æŠ€æœ¯æ”»å…³ï¼ŒæŽ¨åŠ¨æ•°å­—æŠ€æœ¯åˆ›æ–°çªç ´ã€‚
ï¼ˆäºŒï¼‰åŠ é€Ÿæ•°å­—äº§ä¸šåŒ–å‘å±•ã€‚åšå¤§åšå¼ºæ•°å­—äº§ä¸šï¼ŒåŸ¹è‚²å£®å¤§äººå·¥æ™ºèƒ½ã€å¤§æ•°æ®ã€äº‘è®¡ç®—ã€åŒºå—é“¾ç­‰æ–°å…´æ•°å­—äº§ä¸šï¼ŒåŠ å¿«æŽ¨åŠ¨æ•°å­—äº§ä¸šé›†ç¾¤å‘å±•ï¼ŒåŸ¹è‚²å…·æœ‰å›½é™…ç«žäº‰åŠ›çš„æ•°å­—ä¼ä¸šå’Œäº§ä¸šç”Ÿæ€ã€‚
ï¼ˆä¸‰ï¼‰æŽ¨åŠ¨äº§ä¸šæ•°å­—åŒ–è½¬åž‹ã€‚åŠ å¿«ä¼ ç»Ÿäº§ä¸šæ•°å­—åŒ–ã€ç½‘ç»œåŒ–ã€æ™ºèƒ½åŒ–æ”¹é€ ï¼Œæ·±åŒ–æ•°å­—æŠ€æœ¯åœ¨å†œä¸šã€å·¥ä¸šã€æœåŠ¡ä¸šç­‰é¢†åŸŸçš„åº”ç”¨ï¼Œä¿ƒè¿›æ•°å­—æŠ€æœ¯ä¸Žå®žä½“ç»æµŽæ·±åº¦èžåˆï¼ŒåŸ¹è‚²èžåˆæ–°ä¸šæ€ã€æ–°æ¨¡å¼ã€‚
ï¼ˆå››ï¼‰åŠ å¼ºæ•°å­—åŸºç¡€è®¾æ–½å»ºè®¾ã€‚ç³»ç»Ÿå¸ƒå±€æ–°åž‹åŸºç¡€è®¾æ–½ï¼ŒåŠ å¿«5Gç½‘ç»œã€æ•°æ®ä¸­å¿ƒã€å·¥ä¸šäº’è”ç½‘ç­‰æ–°åž‹åŸºç¡€è®¾æ–½å»ºè®¾ï¼Œæž„å»ºé«˜é€Ÿã€ç§»åŠ¨ã€å®‰å…¨ã€æ³›åœ¨çš„æ•°å­—åŸºç¡€è®¾æ–½ä½“ç³»ã€‚
ï¼ˆäº”ï¼‰åŠ å¼ºæ•°æ®èµ„æºå¼€å‘åˆ©ç”¨ã€‚åŠ å¿«æ•°æ®è¦ç´ å¸‚åœºåŒ–ï¼ŒæŽ¨è¿›å…¬å…±æ•°æ®å¼€æ”¾å…±äº«ï¼Œä¿ƒè¿›æ•°æ®èµ„æºæ•´åˆå’Œå¼€å‘åˆ©ç”¨ï¼Œä¿éšœæ•°æ®å®‰å…¨ï¼Œæž„å»ºæ•°æ®è¦ç´ æœ‰åºæµé€šçš„åˆ¶åº¦è§„åˆ™ã€‚
ï¼ˆå…­ï¼‰åŸ¹è‚²æ•°å­—ç»æµŽæ–°ä¸šæ€æ–°æ¨¡å¼ã€‚å¤§åŠ›å‘å±•å¹³å°ç»æµŽã€å…±äº«ç»æµŽã€ç®—æ³•ç»æµŽç­‰æ–°ä¸šæ€ï¼ŒæŽ¨åŠ¨åœ¨çº¿æ•™è‚²ã€è¿œç¨‹åŒ»ç–—ã€æ•°å­—æ–‡åŒ–ç­‰æ–°æ¨¡å¼å‘å±•ï¼Œä¿ƒè¿›çº¿ä¸Šçº¿ä¸‹èžåˆå‘å±•ã€‚
ä¸‰ã€ä¿éšœæŽªæ–½
ï¼ˆä¸€ï¼‰åŠ å¼ºç»„ç»‡é¢†å¯¼å’Œç»Ÿç­¹åè°ƒã€‚å»ºç«‹å¥å…¨å…šå§”é¢†å¯¼ã€æ”¿åºœä¸»å¯¼ã€ä¼ä¸šä¸»ä½“ã€ç¤¾ä¼šå‚ä¸Žçš„æ•°å­—ç»æµŽå‘å±•å·¥ä½œæœºåˆ¶ï¼Œç»Ÿç­¹æŽ¨è¿›æ•°å­—ç»æµŽå‘å±•ã€‚
ï¼ˆäºŒï¼‰å®Œå–„æ•°å­—ç»æµŽæ”¿ç­–ä½“ç³»ã€‚ä¼˜åŒ–æ•°å­—ç»æµŽå‘å±•çš„è´¢ç¨Žã€é‡‘èžã€æŠ•èµ„ã€è´¸æ˜“ç­‰æ”¯æŒæ”¿ç­–ï¼Œå®Œå–„ä¿ƒè¿›æ•°å­—ç»æµŽå‘å±•çš„æ³•å¾‹æ³•è§„å’Œæ ‡å‡†ä½“ç³»ï¼Œæž„å»ºé€‚åº”æ•°å­—ç»æµŽå‘å±•çš„ç›‘ç®¡ä½“ç³»ã€‚
ï¼ˆä¸‰ï¼‰åŠ å¼ºæ•°å­—ç»æµŽäººæ‰åŸ¹å…»ã€‚å®Œå–„æ•°å­—ç»æµŽäººæ‰åŸ¹å…»ä½“ç³»ï¼ŒåŠ å¤§é«˜ç«¯æ•°å­—äººæ‰å¼•è¿›åŠ›åº¦ï¼ŒåŠ å¼ºæ•°å­—æŠ€èƒ½åŸ¹è®­ï¼Œå»ºè®¾é«˜ç´ è´¨æ•°å­—ç»æµŽäººæ‰é˜Ÿä¼ã€‚
ï¼ˆå››ï¼‰å¼ºåŒ–æ•°å­—ç»æµŽå®‰å…¨ä¿éšœã€‚å®Œå–„æ•°å­—å®‰å…¨ä¿éšœä½“ç³»ï¼Œå¼ºåŒ–å…³é”®ä¿¡æ¯åŸºç¡€è®¾æ–½å®‰å…¨ä¿æŠ¤ï¼ŒåŠ å¼ºæ•°æ®å®‰å…¨å’Œä¸ªäººä¿¡æ¯ä¿æŠ¤ï¼Œå¢žå¼ºç½‘ç»œå®‰å…¨é˜²æŠ¤èƒ½åŠ›ã€‚
ï¼ˆäº”ï¼‰æ·±åŒ–æ•°å­—ç»æµŽå›½é™…åˆä½œã€‚ç§‰æŒå¼€æ”¾ã€åˆä½œã€å…±èµ¢ç†å¿µï¼Œç§¯æžå‚ä¸Žæ•°å­—ç»æµŽå›½é™…è§„åˆ™åˆ¶å®šï¼ŒæŽ¨åŠ¨æž„å»ºå¼€æ”¾ã€åŒ…å®¹ã€æ™®æƒ çš„æ•°å­—ç»æµŽå›½é™…åˆä½œä½“ç³»ã€‚
æœ¬æ„è§è‡ªå‘å¸ƒä¹‹æ—¥èµ·æ–½è¡Œã€‚
        """
    }
    
    # åˆ›å»ºä¸´æ—¶ç›®å½•å­˜æ”¾ç¤ºä¾‹æ–‡ä»¶
    example_dir = os.path.join("temp", "examples")
    os.makedirs(example_dir, exist_ok=True)
    
    file_paths = []
    for filename, content in example_files.items():
        file_path = os.path.join(example_dir, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content.strip())
        file_paths.append(file_path)
    
    return file_paths, example_files

def generate_random_text(n_words=500, policy=True):
    """ç”Ÿæˆéšæœºæ–‡æœ¬ç”¨äºŽæµ‹è¯•"""
    # æ”¿ç­–ç›¸å…³è¯æ±‡
    policy_words = ["æ”¿ç­–", "å‘å±•", "æˆ˜ç•¥", "è§„åˆ’", "å®žæ–½", "æŽ¨è¿›", "åˆ›æ–°", "æ”¹é©", "ä¿ƒè¿›",
                    "å»ºè®¾", "åŠ å¼º", "å®Œå–„", "å¥å…¨", "ä¼˜åŒ–", "æå‡", "æ·±åŒ–", "ä¿éšœ", "æ”¯æŒ",
                    "ä½“ç³»", "æœºåˆ¶", "åˆ¶åº¦", "æŽªæ–½", "æ–¹æ¡ˆ", "è¡ŒåŠ¨", "å·¥ç¨‹", "é¡¹ç›®", "è®¡åˆ’"]
    
    # å¸¸ç”¨è¯æ±‡
    common_words = ["çš„", "å’Œ", "åœ¨", "æ˜¯", "äº†", "æœ‰", "ä¸º", "ä»¥", "ä¸Ž", "ä¸", "è¿™", "æˆ‘ä»¬",
                   "è¦", "å¯¹", "ä»Ž", "ç”±", "åˆ°", "ä¸Š", "ä¸­", "ä¸‹", "å„", "ç§", "ç­‰", "å¤š"]
    
    # ç”Ÿæˆéšæœºæ–‡æœ¬
    if policy:
        word_list = policy_words * 3 + common_words
        title = f"éšæœºæ”¿ç­–æ–‡æœ¬_{random.randint(1000, 9999)}"
    else:
        word_list = common_words * 3
        title = f"éšæœºæ–‡æœ¬_{random.randint(1000, 9999)}"
    
    text = title + "\n\n"
    
    # ç”Ÿæˆå‡ ä¸ªæ®µè½
    for i in range(5):
        paragraph = ""
        # æ¯æ®µå¥å­æ•°
        for j in range(random.randint(3, 6)):
            # æ¯å¥è¯æ•°
            sentence_length = random.randint(10, 20)
            sentence = "".join(random.choices(word_list, k=sentence_length))
            paragraph += sentence + "ã€‚"
        text += paragraph + "\n\n"
    
    return text, title

def generate_random_data(n_files=5, show_progress=True):
    """ç”Ÿæˆéšæœºæ•°æ®é›†ç”¨äºŽæµ‹è¯•ï¼ˆæ”¯æŒå¤§è§„æ¨¡ç”Ÿæˆï¼‰"""
    random_dir = os.path.join("temp", "random")
    os.makedirs(random_dir, exist_ok=True)
    
    file_paths = []
    file_contents = {}
    
    # ä½¿ç”¨æ‰¹é‡å¤„ç†æé«˜æ•ˆçŽ‡
    progress_bar = None
    status_text = None
    if show_progress and n_files > 100:
        progress_bar = st.progress(0)
        status_text = st.empty()
    
    for i in range(n_files):
        # æ›´æ–°è¿›åº¦
        if progress_bar and i % 50 == 0:
            progress = (i + 1) / n_files
            progress_bar.progress(progress)
            status_text.text(f"ç”Ÿæˆæ–‡ä»¶ {i+1}/{n_files}")
        
        text, title = generate_random_text(n_words=random.randint(300, 800))
        filename = f"{title}_{i}.txt"  # æ·»åŠ ç´¢å¼•é¿å…é‡å
        file_path = os.path.join(random_dir, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        file_paths.append(file_path)
        file_contents[filename] = text
        
        # æ¯å¤„ç†500ä¸ªæ–‡ä»¶è¿›è¡Œä¸€æ¬¡åžƒåœ¾å›žæ”¶
        if i > 0 and i % 500 == 0:
            gc.collect()
    
    if progress_bar:
        progress_bar.progress(1.0)
        status_text.text("ç”Ÿæˆå®Œæˆ")
    
    return file_paths, file_contents

def detect_encoding(file_path):
    """æ£€æµ‹æ–‡ä»¶ç¼–ç """
    encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'latin-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                f.read()
            return encoding
        except UnicodeDecodeError:
            continue
    return 'utf-8'  # é»˜è®¤è¿”å›žutf-8

def read_file_with_encoding(file_path):
    """ä½¿ç”¨é€‚å½“çš„ç¼–ç è¯»å–æ–‡ä»¶"""
    encoding = detect_encoding(file_path)
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            return f.read()
    except Exception as e:
        log_message(f"è¯»å–æ–‡ä»¶ {os.path.basename(file_path)} å¤±è´¥: {str(e)}", level="error")
        return ""

def read_docx_file(file_path):
    """è¯»å–docxæ–‡ä»¶å†…å®¹"""
    if not DOCX_AVAILABLE:
        log_message(f"æ— æ³•è¯»å– {os.path.basename(file_path)}: python-docxæœªå®‰è£…", level="error")
        return ""
    
    try:
        doc = Document(file_path)
        # æå–æ‰€æœ‰æ®µè½æ–‡æœ¬
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # ä¹Ÿæå–è¡¨æ ¼ä¸­çš„æ–‡æœ¬
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        
        content = '\n'.join(paragraphs)
        return content
    except Exception as e:
        log_message(f"è¯»å–docxæ–‡ä»¶ {os.path.basename(file_path)} å¤±è´¥: {str(e)}", level="error")
        return ""

def read_docx_from_bytes(file_bytes, filename):
    """ä»Žå­—èŠ‚æµè¯»å–docxæ–‡ä»¶å†…å®¹"""
    if not DOCX_AVAILABLE:
        log_message(f"æ— æ³•è¯»å– {filename}: python-docxæœªå®‰è£…", level="error")
        return ""
    
    try:
        import io
        doc = Document(io.BytesIO(file_bytes))
        # æå–æ‰€æœ‰æ®µè½æ–‡æœ¬
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # ä¹Ÿæå–è¡¨æ ¼ä¸­çš„æ–‡æœ¬
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        
        content = '\n'.join(paragraphs)
        return content
    except Exception as e:
        log_message(f"è¯»å–docxæ–‡ä»¶ {filename} å¤±è´¥: {str(e)}", level="error")
        return ""

def get_supported_extensions():
    """èŽ·å–æ”¯æŒçš„æ–‡ä»¶æ‰©å±•ååˆ—è¡¨"""
    extensions = ['.txt']
    if DOCX_AVAILABLE:
        extensions.extend(['.docx', '.doc'])
    return extensions

def is_supported_file(filename):
    """æ£€æŸ¥æ–‡ä»¶æ˜¯å¦ä¸ºæ”¯æŒçš„æ ¼å¼"""
    ext = os.path.splitext(filename.lower())[1]
    return ext in get_supported_extensions()

def read_any_document(file_path):
    """æ ¹æ®æ–‡ä»¶ç±»åž‹è¯»å–æ–‡æ¡£å†…å®¹"""
    ext = os.path.splitext(file_path.lower())[1]
    
    if ext == '.txt':
        return read_file_with_encoding(file_path)
    elif ext in ['.docx', '.doc']:
        return read_docx_file(file_path)
    else:
        log_message(f"ä¸æ”¯æŒçš„æ–‡ä»¶æ ¼å¼: {ext}", level="warning")
        return ""

def render_data_loader():
    """æ¸²æŸ“æ•°æ®åŠ è½½æ¨¡å—"""
    st.header("æ•°æ®åŠ è½½")
    
    # åŠŸèƒ½ä»‹ç»ä¸Žæ“ä½œæ‰‹å†Œ
    with st.expander("ðŸ“– åŠŸèƒ½ä»‹ç»ä¸Žæ“ä½œæ‰‹å†Œ", expanded=False):
        st.markdown("""
        ## ðŸ“‚ æ•°æ®åŠ è½½æ¨¡å—
        
        **åŠŸèƒ½æ¦‚è¿°**ï¼šå¯¼å…¥å¾…åˆ†æžçš„æ”¿ç­–æ–‡æœ¬æ–‡ä»¶ï¼Œæ˜¯LDAä¸»é¢˜æ¨¡åž‹åˆ†æžçš„ç¬¬ä¸€æ­¥ã€‚
        
        ---
        
        ### ðŸŽ¯ ä½¿ç”¨åœºæ™¯
        
        | åœºæ™¯ | æŽ¨èæ–¹å¼ | è¯´æ˜Ž |
        |------|----------|------|
        | åˆæ¬¡ä½“éªŒç³»ç»Ÿ | ç¤ºä¾‹æ•°æ® | å¿«é€Ÿäº†è§£ç³»ç»ŸåŠŸèƒ½ï¼Œæ— éœ€å‡†å¤‡æ•°æ® |
        | æ­£å¼ç ”ç©¶åˆ†æž | ä¸Šä¼ æ–‡ä»¶ | ä¸Šä¼ è‡ªå·±çš„æ”¿ç­–æ–‡æœ¬è¿›è¡Œåˆ†æž |
        | ç³»ç»Ÿæ€§èƒ½æµ‹è¯• | éšæœºæ•°æ® | æµ‹è¯•ç³»ç»Ÿå¤„ç†å¤§è§„æ¨¡æ•°æ®çš„èƒ½åŠ› |
        | å­¦æœ¯è®ºæ–‡ç ”ç©¶ | ä¸Šä¼ æ–‡ä»¶ | ä¸Šä¼ ç ”ç©¶å¯¹è±¡çš„æ”¿ç­–æ–‡æœ¬é›† |
        
        ---
        
        ### ðŸ“‹ æ“ä½œæ­¥éª¤
        
        **æ–¹å¼ä¸€ï¼šä¸Šä¼ æ–‡ä»¶**
        1. é€‰æ‹©"ä¸Šä¼ æ–‡ä»¶"é€‰é¡¹
        2. ç‚¹å‡»ä¸Šä¼ åŒºåŸŸï¼Œé€‰æ‹©TXTæ–‡ä»¶æˆ–ZIPåŽ‹ç¼©åŒ…
        3. ç‚¹å‡»"å¼€å§‹åŠ è½½"æŒ‰é’®
        4. ç­‰å¾…æ–‡ä»¶å¤„ç†å®Œæˆ
        
        **æ–¹å¼äºŒï¼šä½¿ç”¨ç¤ºä¾‹æ•°æ®**
        1. é€‰æ‹©"ä½¿ç”¨ç¤ºä¾‹æ•°æ®"é€‰é¡¹
        2. ç‚¹å‡»"åŠ è½½ç¤ºä¾‹æ•°æ®"æŒ‰é’®
        3. ç³»ç»Ÿå°†åŠ è½½3ä»½é¢„ç½®çš„æ”¿ç­–æ–‡ä»¶
        
        **æ–¹å¼ä¸‰ï¼šç”Ÿæˆéšæœºæ•°æ®**
        1. é€‰æ‹©"ç”Ÿæˆéšæœºæ•°æ®"é€‰é¡¹
        2. è®¾ç½®è¦ç”Ÿæˆçš„æ–‡ä»¶æ•°é‡ï¼ˆæˆ–ä½¿ç”¨å¿«æ·é€‰é¡¹ï¼‰
        3. ç‚¹å‡»"ç”Ÿæˆéšæœºæ•°æ®"æŒ‰é’®
        
        ---
        
        ### ðŸ“ æ”¯æŒçš„æ–‡ä»¶æ ¼å¼
        
        | æ ¼å¼ | è¯´æ˜Ž | ç¼–ç æ”¯æŒ |
        |------|------|----------|
        | .txt | çº¯æ–‡æœ¬æ–‡ä»¶ | UTF-8, GBK, GB2312, GB18030, Big5 |
        | .docx | Wordæ–‡æ¡£ï¼ˆæŽ¨èï¼‰ | è‡ªåŠ¨è§£æž |
        | .doc | æ—§ç‰ˆWordæ–‡æ¡£ | è‡ªåŠ¨è§£æž |
        | .zip | åŽ‹ç¼©åŒ…ï¼ˆå†…å«å¤šä¸ªæ–‡æ¡£ï¼‰ | è‡ªåŠ¨æ£€æµ‹ |
        
        **æ³¨æ„**ï¼šDOC/DOCXæ”¯æŒéœ€è¦å®‰è£…python-docxåº“ã€‚ZIPåŽ‹ç¼©åŒ…å¯åŒ…å«TXTã€DOCã€DOCXæ··åˆæ–‡ä»¶ã€‚
        
        ---
        
        ### âš™ï¸ å‚æ•°è¯´æ˜Ž
        
        | å‚æ•° | èŒƒå›´ | é»˜è®¤å€¼ | è¯´æ˜Ž |
        |------|------|--------|------|
        | æœ€å¤§æ–‡ä»¶æ•° | 1-10000 | 10000 | ç³»ç»Ÿæ”¯æŒçš„æœ€å¤§æ–‡ä»¶æ•°é‡ |
        | æ‰¹é‡å¤„ç†å¤§å° | - | 100 | æ¯æ‰¹å¤„ç†çš„æ–‡ä»¶æ•°é‡ |
        | å†…å­˜è­¦å‘Šé˜ˆå€¼ | - | 1000 | è¶…è¿‡æ­¤æ•°é‡æ˜¾ç¤ºå†…å­˜è­¦å‘Š |
        
        ---
        
        ### ðŸ’¡ ä½¿ç”¨å»ºè®®
        
        **æ•°æ®å‡†å¤‡å»ºè®®ï¼š**
        - æ¯ä¸ªTXTæ–‡ä»¶åŒ…å«ä¸€ä»½å®Œæ•´çš„æ”¿ç­–æ–‡æ¡£
        - æ–‡ä»¶åå»ºè®®åŒ…å«æ”¿ç­–åç§°æˆ–ç¼–å·ï¼Œä¾¿äºŽåŽç»­è¯†åˆ«
        - ç¡®ä¿æ–‡æœ¬å†…å®¹ä¸ºçº¯æ–‡æœ¬æ ¼å¼ï¼Œé¿å…åŒ…å«ç‰¹æ®Šæ ¼å¼
        
        **å¤§è§„æ¨¡æ•°æ®å¤„ç†ï¼š**
        - è¶…è¿‡100ä¸ªæ–‡ä»¶å»ºè®®ä½¿ç”¨ZIPåŽ‹ç¼©åŒ…ä¸Šä¼ 
        - è¶…è¿‡1000ä¸ªæ–‡ä»¶æ—¶æ³¨æ„å†…å­˜ä½¿ç”¨
        - å¤„ç†å®ŒæˆåŽå¯ç‚¹å‡»"æ¸…ç©ºæ•°æ®"é‡Šæ”¾å†…å­˜
        
        **å­¦æœ¯ç ”ç©¶å»ºè®®ï¼š**
        - å»ºè®®æ”¶é›†åŒä¸€é¢†åŸŸã€åŒä¸€æ—¶æœŸçš„æ”¿ç­–æ–‡æœ¬
        - æ–‡æœ¬æ•°é‡å»ºè®®åœ¨50-500ä»½ä¹‹é—´
        - æ¯ä»½æ–‡æœ¬é•¿åº¦å»ºè®®åœ¨500-5000å­—ä¹‹é—´
        
        ---
        
        ### â“ å¸¸è§é—®é¢˜
        
        **Q: æ–‡ä»¶ä¸Šä¼ å¤±è´¥æ€Žä¹ˆåŠžï¼Ÿ**
        A: æ£€æŸ¥æ–‡ä»¶ç¼–ç æ˜¯å¦ä¸ºæ”¯æŒçš„æ ¼å¼ï¼Œå°è¯•å°†æ–‡ä»¶è½¬æ¢ä¸ºUTF-8ç¼–ç ã€‚
        
        **Q: ç³»ç»Ÿæ˜¾ç¤ºå†…å­˜è­¦å‘Šæ€Žä¹ˆåŠžï¼Ÿ**
        A: å¯ä»¥å‡å°‘åŠ è½½çš„æ–‡ä»¶æ•°é‡ï¼Œæˆ–åœ¨åˆ†æžå®ŒæˆåŽç‚¹å‡»"æ¸…ç©ºæ•°æ®"é‡Šæ”¾å†…å­˜ã€‚
        
        **Q: å¦‚ä½•åˆ¤æ–­æ•°æ®åŠ è½½æ˜¯å¦æˆåŠŸï¼Ÿ**
        A: åŠ è½½æˆåŠŸåŽä¼šæ˜¾ç¤ºæ–‡ä»¶åˆ—è¡¨å’Œç»Ÿè®¡ä¿¡æ¯ï¼Œå¯ä»¥é¢„è§ˆæ–‡ä»¶å†…å®¹ç¡®è®¤ã€‚
        """)
    
    # æ•°æ®åŠ è½½é€‰é¡¹
    data_source = st.radio(
        "é€‰æ‹©æ•°æ®æ¥æº",
        ["ä¸Šä¼ æ–‡ä»¶", "ä½¿ç”¨ç¤ºä¾‹æ•°æ®", "ç”Ÿæˆéšæœºæ•°æ®"],
        horizontal=True,
        key="data_source_radio"
    )
    
    if data_source == "ä¸Šä¼ æ–‡ä»¶":
        st.subheader("ä¸Šä¼ æ”¿ç­–æ–‡ä»¶")
        
        # æ ¹æ®æ˜¯å¦å®‰è£…python-docxæ˜¾ç¤ºä¸åŒæç¤º
        if DOCX_AVAILABLE:
            st.info("ðŸ’¡ æç¤ºï¼šæ”¯æŒTXTã€DOCã€DOCXæ–‡ä»¶ï¼Œå¤§é‡æ–‡ä»¶å»ºè®®æ‰“åŒ…æˆZIPåŽ‹ç¼©åŒ…ä¸Šä¼ ï¼Œæ”¯æŒæœ€å¤š10000ä¸ªæ–‡ä»¶")
            accepted_types = ["txt", "doc", "docx", "zip"]
        else:
            st.info("ðŸ’¡ æç¤ºï¼šå¤§é‡æ–‡ä»¶å»ºè®®æ‰“åŒ…æˆZIPåŽ‹ç¼©åŒ…ä¸Šä¼ ï¼Œæ”¯æŒæœ€å¤š10000ä¸ªæ–‡ä»¶")
            st.warning("âš ï¸ python-docxæœªå®‰è£…ï¼Œæš‚ä¸æ”¯æŒDOC/DOCXæ–‡ä»¶ã€‚å¦‚éœ€æ”¯æŒè¯·å®‰è£…: pip install python-docx")
            accepted_types = ["txt", "zip"]
        
        uploaded_files = st.file_uploader(
            "ä¸Šä¼ æ–‡æœ¬æ–‡ä»¶æˆ–ZIPåŽ‹ç¼©åŒ…", 
            type=accepted_types, 
            accept_multiple_files=True,
            key="data_files_uploader"
        )
        
        if uploaded_files and st.button("å¼€å§‹åŠ è½½", key="load_uploaded_files"):
            with st.spinner("æ­£åœ¨å¤„ç†ä¸Šä¼ æ–‡ä»¶..."):
                st.session_state.file_contents = {}
                st.session_state.file_names = []
                st.session_state.raw_texts = []
                
                total_files = len(uploaded_files)
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                processed_count = 0
                skipped_count = 0
                
                for i, uploaded_file in enumerate(uploaded_files):
                    status_text.text(f"å¤„ç†æ–‡ä»¶ {i+1}/{total_files}: {uploaded_file.name}")
                    
                    if uploaded_file.name.endswith('.zip'):
                        # å¤„ç†ZIPæ–‡ä»¶
                        with tempfile.TemporaryDirectory() as temp_dir:
                            zip_path = os.path.join(temp_dir, uploaded_file.name)
                            with open(zip_path, 'wb') as f:
                                f.write(uploaded_file.getbuffer())
                            
                            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                                # èŽ·å–ZIPä¸­æ”¯æŒçš„æ–‡ä»¶åˆ—è¡¨
                                supported_files = [f for f in zip_ref.namelist() if is_supported_file(f)]
                                total_supported = len(supported_files)
                                
                                zip_ref.extractall(temp_dir)
                                
                                # æ‰¹é‡å¤„ç†ZIPä¸­çš„æ–‡ä»¶
                                for j, file in enumerate(supported_files):
                                    if j % 50 == 0:
                                        progress = (i + j/max(total_supported, 1)) / total_files
                                        progress_bar.progress(min(progress, 1.0))
                                        status_text.text(f"å¤„ç†ZIPæ–‡ä»¶ {i+1}/{total_files}ï¼Œå†…éƒ¨æ–‡ä»¶ {j+1}/{total_supported}")
                                    
                                    file_path = os.path.join(temp_dir, file)
                                    if os.path.isfile(file_path):
                                        content = read_any_document(file_path)
                                        if content and content.strip():
                                            filename = os.path.basename(file)
                                            st.session_state.file_contents[filename] = content
                                            st.session_state.file_names.append(filename)
                                            st.session_state.raw_texts.append(content)
                                            processed_count += 1
                                        else:
                                            skipped_count += 1
                                    
                                    # å®šæœŸåžƒåœ¾å›žæ”¶
                                    if j > 0 and j % 500 == 0:
                                        gc.collect()
                    
                    elif uploaded_file.name.lower().endswith('.txt'):
                        # å¤„ç†å•ä¸ªTXTæ–‡ä»¶
                        content = uploaded_file.read().decode('utf-8', errors='replace')
                        if content and content.strip():
                            st.session_state.file_contents[uploaded_file.name] = content
                            st.session_state.file_names.append(uploaded_file.name)
                            st.session_state.raw_texts.append(content)
                            processed_count += 1
                        else:
                            skipped_count += 1
                    
                    elif uploaded_file.name.lower().endswith(('.docx', '.doc')):
                        # å¤„ç†DOC/DOCXæ–‡ä»¶
                        if DOCX_AVAILABLE:
                            file_bytes = uploaded_file.read()
                            content = read_docx_from_bytes(file_bytes, uploaded_file.name)
                            if content and content.strip():
                                st.session_state.file_contents[uploaded_file.name] = content
                                st.session_state.file_names.append(uploaded_file.name)
                                st.session_state.raw_texts.append(content)
                                processed_count += 1
                            else:
                                skipped_count += 1
                                log_message(f"æ–‡ä»¶ {uploaded_file.name} å†…å®¹ä¸ºç©ºï¼Œå·²è·³è¿‡", level="warning")
                        else:
                            skipped_count += 1
                            log_message(f"æ— æ³•å¤„ç† {uploaded_file.name}: python-docxæœªå®‰è£…", level="warning")
                    
                    progress_bar.progress((i + 1) / total_files)
                
                progress_bar.progress(1.0)
                status_text.text("æ–‡ä»¶åŠ è½½å®Œæˆ")
                update_progress(1.0, "æ–‡ä»¶åŠ è½½å®Œæˆ")
                log_message(f"å·²åŠ è½½ {processed_count} ä¸ªæ–‡ä»¶", level="success")
                
                # æ˜¾ç¤ºåŠ è½½ç»“æžœ
                if skipped_count > 0:
                    st.success(f"æˆåŠŸåŠ è½½ {processed_count} ä¸ªæ–‡ä»¶ï¼Œè·³è¿‡ {skipped_count} ä¸ªç©ºæ–‡ä»¶æˆ–ä¸æ”¯æŒçš„æ–‡ä»¶")
                else:
                    st.success(f"æˆåŠŸåŠ è½½ {processed_count} ä¸ªæ–‡ä»¶")
                
                # å†…å­˜è­¦å‘Š
                if processed_count > MEMORY_WARNING_THRESHOLD:
                    st.warning(f"âš ï¸ å·²åŠ è½½ {processed_count} ä¸ªæ–‡ä»¶ï¼Œè¯·æ³¨æ„å†…å­˜ä½¿ç”¨ã€‚å»ºè®®åœ¨å¤„ç†å®ŒæˆåŽæ¸…ç†ç¼“å­˜ã€‚")
                
                # æ¸…ç©ºä¸Šä¼ çš„æ–‡ä»¶
                st.session_state.uploaded_files = None
    
    elif data_source == "ä½¿ç”¨ç¤ºä¾‹æ•°æ®":
        st.subheader("ç¤ºä¾‹æ”¿ç­–æ•°æ®")
        st.info("è¿™äº›æ˜¯é¢„è®¾çš„æ”¿ç­–æ–‡ä»¶ç¤ºä¾‹ï¼Œç”¨äºŽç³»ç»ŸåŠŸèƒ½æ¼”ç¤ºã€‚")
        
        if st.button("åŠ è½½ç¤ºä¾‹æ•°æ®", key="load_example_data"):
            with st.spinner("æ­£åœ¨åŠ è½½ç¤ºä¾‹æ•°æ®..."):
                file_paths, file_contents = load_example_data()
                
                st.session_state.file_contents = file_contents
                st.session_state.file_names = list(file_contents.keys())
                st.session_state.raw_texts = list(file_contents.values())
                
                log_message(f"å·²åŠ è½½ {len(file_paths)} ä¸ªç¤ºä¾‹æ–‡ä»¶", level="success")
                st.success(f"æˆåŠŸåŠ è½½ {len(file_paths)} ä¸ªç¤ºä¾‹æ–‡ä»¶")
    
    elif data_source == "ç”Ÿæˆéšæœºæ•°æ®":
        st.subheader("ç”Ÿæˆéšæœºæµ‹è¯•æ•°æ®")
        
        # æ‰©å±•æ»‘å—èŒƒå›´ä»¥æ”¯æŒå¤§è§„æ¨¡æµ‹è¯•
        col1, col2 = st.columns(2)
        with col1:
            n_files = st.number_input(
                "ç”Ÿæˆæ–‡ä»¶æ•°é‡", 
                min_value=3, 
                max_value=MAX_FILES, 
                value=100,
                step=100,
                help=f"æ”¯æŒç”Ÿæˆ3-{MAX_FILES}ä¸ªæµ‹è¯•æ–‡ä»¶",
                key="n_files_input"
            )
        with col2:
            # å¿«æ·é€‰é¡¹
            quick_options = st.selectbox(
                "å¿«æ·é€‰é¡¹",
                ["è‡ªå®šä¹‰", "100ä¸ª(å¿«é€Ÿæµ‹è¯•)", "500ä¸ª(ä¸­ç­‰è§„æ¨¡)", "2000ä¸ª(å¤§è§„æ¨¡)", "5000ä¸ª(åŽ‹åŠ›æµ‹è¯•)"],
                key="quick_file_options"
            )
            if quick_options != "è‡ªå®šä¹‰":
                n_files = int(quick_options.split("ä¸ª")[0])
        
        # å¤§è§„æ¨¡ç”Ÿæˆè­¦å‘Š
        if n_files > 1000:
            st.warning(f"âš ï¸ å³å°†ç”Ÿæˆ {n_files} ä¸ªæ–‡ä»¶ï¼Œè¿™å¯èƒ½éœ€è¦è¾ƒé•¿æ—¶é—´å’Œè¾ƒå¤šå†…å­˜ã€‚")
        
        if st.button("ç”Ÿæˆéšæœºæ•°æ®", key="generate_random_data"):
            start_time = time.time()
            
            with st.spinner(f"æ­£åœ¨ç”Ÿæˆ {n_files} ä¸ªéšæœºæ–‡ä»¶..."):
                file_paths, file_contents = generate_random_data(n_files, show_progress=True)
                
                st.session_state.file_contents = file_contents
                st.session_state.file_names = list(file_contents.keys())
                st.session_state.raw_texts = list(file_contents.values())
                
                elapsed = time.time() - start_time
                log_message(f"å·²ç”Ÿæˆ {len(file_paths)} ä¸ªéšæœºæ–‡ä»¶ï¼Œè€—æ—¶ {elapsed:.1f}ç§’", level="success")
                st.success(f"æˆåŠŸç”Ÿæˆ {len(file_paths)} ä¸ªéšæœºæ–‡ä»¶ï¼Œè€—æ—¶ {elapsed:.1f}ç§’")
                
                # åžƒåœ¾å›žæ”¶
                gc.collect()
    
    # æ˜¾ç¤ºå·²åŠ è½½çš„æ–‡ä»¶
    if st.session_state.get("file_contents"):
        st.subheader("å·²åŠ è½½çš„æ–‡ä»¶")
        
        # æ·»åŠ æ¸…ç©ºæ•°æ®æŒ‰é’®
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            file_count = len(st.session_state.file_names)
            total_chars = sum(len(text) for text in st.session_state.raw_texts)
            st.metric("æ–‡ä»¶ç»Ÿè®¡", f"{file_count} ä¸ªæ–‡ä»¶ï¼Œå…± {total_chars:,} å­—ç¬¦")
        with col3:
            if st.button("ðŸ—‘ï¸ æ¸…ç©ºæ•°æ®", key="clear_loaded_data", type="secondary"):
                st.session_state.file_contents = {}
                st.session_state.file_names = []
                st.session_state.raw_texts = []
                st.session_state.texts = None
                st.session_state.dictionary = None
                st.session_state.corpus = None
                st.session_state.lda_model = None
                st.session_state.training_complete = False
                st.session_state.pyldavis_html = None
                st.session_state.wordcloud_images = {}
                gc.collect()  # å¼ºåˆ¶åžƒåœ¾å›žæ”¶
                log_message("å·²æ¸…ç©ºæ‰€æœ‰åŠ è½½çš„æ•°æ®", level="warning")
                st.rerun()
        
        # å¤§é‡æ–‡ä»¶æ—¶åªæ˜¾ç¤ºæ‘˜è¦
        file_count = len(st.session_state.file_names)
        if file_count > 100:
            st.info(f"ðŸ“Š å·²åŠ è½½ {file_count} ä¸ªæ–‡ä»¶ï¼ˆæ•°æ®é‡è¾ƒå¤§ï¼Œä»…æ˜¾ç¤ºå‰100ä¸ªï¼‰")
            display_names = st.session_state.file_names[:100]
            display_texts = st.session_state.raw_texts[:100]
        else:
            display_names = st.session_state.file_names
            display_texts = st.session_state.raw_texts
        
        df = pd.DataFrame({
            "æ–‡ä»¶å": display_names,
            "æ–‡æœ¬é•¿åº¦": [len(text) for text in display_texts]
        })
        
        st.dataframe(df, use_container_width=True, height=300)
        
        # æ–‡ä»¶é¢„è§ˆ
        if st.session_state.file_names:
            with st.expander("æ–‡ä»¶å†…å®¹é¢„è§ˆ", expanded=False):
                # é™åˆ¶é¢„è§ˆé€‰é¡¹æ•°é‡
                preview_options = st.session_state.file_names[:100] if file_count > 100 else st.session_state.file_names
                preview_file = st.selectbox("é€‰æ‹©æ–‡ä»¶é¢„è§ˆ", preview_options, key="data_preview_file_select")
                if preview_file:
                    content = st.session_state.file_contents[preview_file]
                    # é™åˆ¶é¢„è§ˆå†…å®¹é•¿åº¦
                    max_preview_length = 1000
                    if len(content) > max_preview_length:
                        preview_content = content[:max_preview_length] + "..."
                    else:
                        preview_content = content
                    
                    st.text_area("æ–‡ä»¶å†…å®¹", preview_content, height=300) 